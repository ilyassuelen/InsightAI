from __future__ import annotations

import tests as _test_bootstrap  # noqa: F401  # configure isolated services first
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document as WordDocument
from docling.datamodel.base_models import InputFormat

from backend.parsers.docx_parser import parse_docx
from backend.parsers.pdf_parser import create_converter, pdf_contains_text
from backend.parsers.txt_parser import parse_txt


class TextParserTests(unittest.TestCase):
    def test_txt_parser_reads_utf8_text(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.txt"
            path.write_text("InsightAI äöü", encoding="utf-8")
            self.assertEqual(parse_txt(path), "InsightAI äöü")

    def test_txt_parser_ignores_invalid_utf8_bytes(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "invalid.txt"
            path.write_bytes(b"before\xffafter")
            self.assertEqual(parse_txt(path), "beforeafter")

    def test_txt_parser_rejects_missing_file(self) -> None:
        with self.assertRaises(FileNotFoundError):
            parse_txt("/tmp/insightai-definitely-missing.txt")


class DocxParserTests(unittest.TestCase):
    def test_docx_parser_returns_non_empty_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.docx"
            document = WordDocument()
            document.add_paragraph("First paragraph")
            document.add_paragraph("   ")
            document.add_paragraph("Second paragraph")
            document.save(path)

            self.assertEqual(parse_docx(path), "First paragraph\n\nSecond paragraph")

    def test_docx_parser_preserves_headings_lists_tables_and_order(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "structured.docx"
            document = WordDocument()
            document.add_heading("Overview", level=1)
            document.add_paragraph("Introductory paragraph")
            document.add_paragraph("First item", style="List Bullet")
            document.add_paragraph("Nested item", style="List Bullet 2")
            document.add_paragraph("Numbered item", style="List Number")

            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "A | B"
            table.cell(1, 1).text = "10"

            document.add_heading("Details", level=2)
            document.add_paragraph("Closing paragraph")
            document.save(path)

            self.assertEqual(
                parse_docx(path),
                (
                    "# Overview\n\n"
                    "Introductory paragraph\n\n"
                    "- First item\n"
                    "    - Nested item\n"
                    "1. Numbered item\n\n"
                    "| Name | Value |\n"
                    "| --- | --- |\n"
                    "| A \\| B | 10 |\n\n"
                    "## Details\n\n"
                    "Closing paragraph"
                ),
            )

    def test_docx_parser_supports_custom_outline_heading_styles(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "custom-heading.docx"
            document = WordDocument()
            style = document.styles.add_style("Insight Section", 1)
            style.paragraph_format.keep_with_next = True
            style.element.get_or_add_pPr().get_or_add_outlineLvl().val = 2
            document.add_paragraph("Custom section", style=style)
            document.add_paragraph("Section content")
            document.save(path)

            self.assertEqual(
                parse_docx(path),
                "### Custom section\n\nSection content",
            )

    def test_docx_parser_rejects_missing_file(self) -> None:
        with self.assertRaises(FileNotFoundError):
            parse_docx("/tmp/insightai-definitely-missing.docx")


class PdfParserTests(unittest.TestCase):
    def _create_pdf(self, path: Path, text: str | None) -> None:
        document = fitz.open()
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
        document.save(path)
        document.close()

    def test_pdf_text_detection_distinguishes_text_and_scan_like_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            text_pdf = Path(directory) / "text.pdf"
            blank_pdf = Path(directory) / "blank.pdf"
            self._create_pdf(text_pdf, "Extractable text")
            self._create_pdf(blank_pdf, None)

            self.assertTrue(pdf_contains_text(str(text_pdf)))
            self.assertFalse(pdf_contains_text(str(blank_pdf)))

    def test_pdf_detection_failure_defaults_to_ocr_needed(self) -> None:
        self.assertFalse(pdf_contains_text("/tmp/not-a-real-pdf.pdf"))

    def test_converter_disables_expensive_features_and_selects_ocr(self) -> None:
        with patch("backend.parsers.pdf_parser.pdf_contains_text", return_value=True):
            converter = create_converter("unused.pdf")
            options = converter.format_to_options[InputFormat.PDF].pipeline_options
            self.assertFalse(options.do_ocr)
            self.assertFalse(options.do_table_structure)
            self.assertFalse(options.do_picture_description)

        with patch("backend.parsers.pdf_parser.pdf_contains_text", return_value=False):
            converter = create_converter("unused.pdf")
            options = converter.format_to_options[InputFormat.PDF].pipeline_options
            self.assertTrue(options.do_ocr)


if __name__ == "__main__":
    unittest.main()
